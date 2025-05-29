import cv2
import numpy as np
import os
import base64
from flask import current_app
from sqlalchemy import func
from transliterate import translit
import qrcode
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime


def generate_post_address(header):
    address = translit(header, reversed=True).lower().replace(' ', '_')
    return address


def get_next_filename(folder, extension):
    if not os.path.exists(folder):
        os.makedirs(folder)
        return f'1.{extension}'

    existing_files = [f for f in os.listdir(folder) if f.endswith(f".{extension}")]
    numbers = []

    for file in existing_files:
        try:
            num = int(file.split('.')[0])
            numbers.append(num)
        except ValueError:
            continue

    next_number = max(numbers, default=0) + 1
    return f'{next_number}.{extension}'

def enhance_and_resize(img, target_width=1920, target_height=1080):

    img = cv2.resize(img, (target_width, target_height), interpolation=cv2.INTER_LANCZOS4)

    return img


def save_image(image, post_address, image_type):
    try:
        image_folder = f'src/assets/{post_address}/{image_type}'
        os.makedirs(image_folder, exist_ok=True)

        image_format = image.split(";")[0].split("/")[1]
        base64_data = image.split(",")[1]
        image_data = base64.b64decode(base64_data)

        np_arr = np.frombuffer(image_data, np.uint8)
        img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)

        if img is None:
            raise ValueError("Ошибка декодирования изображения")

        img = enhance_and_resize(img)

        filename = get_next_filename(image_folder, image_format)
        image_path = os.path.join(image_folder, filename)

        if not allowed_file(filename):
            raise ValueError(f'Недопустимое расширение файла: {filename}')

        cv2.imwrite(image_path, img)
        print(f"Изображение сохранено: {image_path}")

        return image_path

    except Exception as e:
        print(f"Ошибка при обработке изображения: {e}")



def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_IMAGE_EXTENSIONS']


def convert_json_date_to_sqlite_format(json_date, date_key):
    return func.date(
        func.substr(func.json_extract(json_date, date_key), 7, 4) + '-' +
        func.substr(func.json_extract(json_date, date_key), 4, 2) + '-' +
        func.substr(func.json_extract(json_date, date_key), 1, 2)
    )


def generate_qr_code(address):

    BASE_ADDRESS = 'http://localhost:5173/post'

    try:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )

        full_address = f'{BASE_ADDRESS}/{address}'

        qr.add_data(full_address)
        qr.make(fit=True)

        img = qr.make_image(fill_color="black", back_color="white")

        return img
    except Exception as e:
        raise ValueError(f'Ошибка при генерации QR-кода: {e}')


def parse_post_dates(post):
    try:
        if not post.date_range:
            return None, None

        date_data = json.loads(post.date_range)
        start_date_str = str(date_data.get('start_date', '')) if date_data.get('start_date') else ''
        end_date_str = str(date_data.get('end_date', '')) if date_data.get('end_date') else ''

        post_start = None
        if start_date_str:
            if len(start_date_str) == 4 and start_date_str.isdigit():
                post_start = datetime(int(start_date_str), 1, 1).date()
            else:
                try:
                    post_start = datetime.strptime(start_date_str, "%Y-%m-%d").date()
                except ValueError:
                    pass

        post_end = None
        if end_date_str:
            if len(end_date_str) == 4 and end_date_str.isdigit():
                post_end = datetime(int(end_date_str), 12, 31).date()
            else:
                try:
                    post_end = datetime.strptime(end_date_str, "%Y-%m-%d").date()
                except ValueError:
                    pass

        return post_start, post_end
    except:
        return None, None

def generate_doc_with_qr_bytes(header, address):
    try:
        qr_img = generate_qr_code(address)
        qr_bytes = BytesIO()
        qr_img.save(qr_bytes, format='PNG')
        qr_bytes.seek(0)

        document = Document()

        logo_path = './src/static/logo.png'
        if os.path.exists(logo_path):
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(6))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = document.add_paragraph()
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(20)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        qr_width_in_inches = 6
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(qr_bytes, width=Inches(qr_width_in_inches))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_bytes = BytesIO()
        document.save(doc_bytes)
        doc_bytes.seek(0)

        return doc_bytes
    except Exception as e:
        raise ValueError(f'Ошибка при создании Word-документа: {e}')
